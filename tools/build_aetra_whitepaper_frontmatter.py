from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("docs/whitepaper/Aetra_Whitepaper_Frontmatter_Draft.docx")
BODY_FONT = "Latin Modern Roman"
ACCENT_FONT = "Computer Modern Unicode"
FALLBACK_FONT = "Times New Roman"


def set_font(run, size=None, bold=False, italic=False, small_caps=False, spacing=None):
    font = ACCENT_FONT if small_caps else BODY_FONT
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
    run.bold = bold
    run.italic = italic
    run.font.small_caps = small_caps
    if size is not None:
        run.font.size = Pt(size)
    if spacing is not None:
        rpr = run._element.get_or_add_rPr()
        sp = rpr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            rpr.append(sp)
        sp.set(qn("w:val"), str(spacing))


def set_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.42)
    section.right_margin = Inches(1.42)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = True
    sect_pr = section._sectPr
    title_pg = sect_pr.find(qn("w:titlePg"))
    if title_pg is None:
        title_pg = OxmlElement("w:titlePg")
        sect_pr.append(title_pg)


def bottom_rule(paragraph, color="6B6B6B", size="6"):
    ppr = paragraph._p.get_or_add_pPr()
    bdr = ppr.find(qn("w:pBdr"))
    if bdr is None:
        bdr = OxmlElement("w:pBdr")
        ppr.append(bdr)
    bottom = bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run, 9)
    for kind, text in (
        ("begin", None),
        (None, " PAGE "),
        ("separate", None),
        (None, "1"),
        ("end", None),
    ):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        elif text == " PAGE ":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = text
        else:
            node = OxmlElement("w:t")
            node.text = text
        run._r.append(node)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.03
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = styles["Heading 1"]
    h1.font.name = ACCENT_FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), ACCENT_FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), ACCENT_FONT)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = None
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(9)

    h2 = styles["Heading 2"]
    h2.font.name = ACCENT_FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), ACCENT_FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), ACCENT_FONT)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = None
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)

    if "AbstractBody" not in styles:
        s = styles.add_style("AbstractBody", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = BODY_FONT
        s._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
        s.font.size = Pt(10.5)
        s.paragraph_format.left_indent = Inches(0.45)
        s.paragraph_format.right_indent = Inches(0.45)
        s.paragraph_format.line_spacing = 1.02
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "ContentsMajor" not in styles:
        s = styles.add_style("ContentsMajor", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = ACCENT_FONT
        s._element.rPr.rFonts.set(qn("w:ascii"), ACCENT_FONT)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), ACCENT_FONT)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
        s.font.size = Pt(12)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(7)
        s.paragraph_format.space_after = Pt(3)

    if "ContentsMinor" not in styles:
        s = styles.add_style("ContentsMinor", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = BODY_FONT
        s._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_FONT)
        s.font.size = Pt(10.5)
        s.font.bold = False
        s.paragraph_format.left_indent = Inches(0.27)
        s.paragraph_format.space_after = Pt(1)


def paragraph(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_font(r)
    if align is not None:
        p.alignment = align
    return p


def add_run(p, text, **kwargs):
    r = p.add_run(text)
    set_font(r, **kwargs)
    return r


def set_running_header(section, text):
    section.different_first_page_header_footer = True
    sect_pr = section._sectPr
    title_pg = sect_pr.find(qn("w:titlePg"))
    if title_pg is None:
        title_pg = OxmlElement("w:titlePg")
        sect_pr.append(title_pg)

    first = section.first_page_header.paragraphs[0]
    first.text = ""
    first.paragraph_format.space_after = Pt(0)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(1)
    add_run(hp, text, size=9, small_caps=True, spacing=18)
    bottom_rule(hp)
    add_page_number(section.first_page_footer.paragraphs[0])
    add_page_number(section.footer.paragraphs[0])


def clear_section_links(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False


def title_page(doc):
    for _ in range(3):
        paragraph(doc)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(9)
    add_run(p, "Aetra", size=21, spacing=34)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(14)
    add_run(p, "decentralized proof-of-stake execution network", size=14, spacing=8)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "Daniil Shcherbakov", size=12, spacing=4)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(20)
    add_run(p, "July 8, 2026", size=12, spacing=4)

    for _ in range(2):
        paragraph(doc)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Abstract", size=10.5, bold=True)

    paragraph(
        doc,
        "The aim of this text is to provide a first outline of Aetra: a "
        "decentralized blockchain network designed for trust, moderate speed, "
        "and practical validator operation rather than maximum throughput at "
        "any cost. Aetra targets stronger-than-average decentralization, a "
        "medium hardware profile, bounded validator influence, and native "
        "smart contracts based on the Aetra Virtual Machine.",
        style="AbstractBody",
    )
    paragraph(
        doc,
        "The system is intended to support proof-of-stake validation, nominator "
        "pools, controlled inflation, storage accountability, and future "
        "throughput scaling without making synchronization or validator entry "
        "unnecessarily hard.",
        style="AbstractBody",
    )

    paragraph(doc)
    h = paragraph(doc)
    add_run(h, "Introduction", size=16, bold=True, spacing=12)

    p = paragraph(doc)
    add_run(p, "Aetra", italic=True)
    add_run(
        p,
        " is a blockchain network project built around a conservative "
        "trade-off: it should be faster and more usable than slow settlement "
        "layers, but it should not sacrifice decentralization, auditability, "
        "or validator accessibility in order to compete for the highest "
        "possible transaction count.",
    )

    paragraph(
        doc,
        "The validator economy is an important part of this design. Aetra "
        "should not let the largest holders grow faster only because they "
        "already control more stake. Validator power caps, nominator pools, "
        "commission bounds, concentration metrics, and reward modifiers are "
        "intended to reduce cartel formation and pressure against a small set "
        "of operators.",
    )

    paragraph(
        doc,
        "At the same time, the network must remain economically realistic. "
        "Validators should earn enough to justify reliable infrastructure, "
        "monitoring, and operational risk, but the protocol should avoid "
        "turning high yield into its main product.",
    )

    paragraph(
        doc,
        "Aetra also includes a native virtual machine for smart contracts. "
        "Application logic such as tokens, NFTs, domains, markets, and exchange "
        "contracts should be implemented through AVM standards, while "
        "protocol-critical systems remain native parts of the chain.",
    )


def contents_page(doc):
    doc.add_page_break()
    h = paragraph(doc)
    add_run(h, "Contents", size=16, bold=True, spacing=10)

    items = [
        ("1", "Brief Description of Aetra Components", [
            ("1.1", "Aether Core"),
            ("1.2", "Validator Economy"),
            ("1.3", "Aetra Virtual Machine"),
            ("1.4", "System Entities"),
        ]),
        ("2", "Aetra Blockchain", [
            ("2.1", "Accounts, Addresses, and State"),
            ("2.2", "Messages, Transactions, and Receipts"),
            ("2.3", "Genesis, Export, and Import"),
            ("2.4", "Upgrades and Invariants"),
        ]),
        ("3", "Consensus and Staking", [
            ("3.1", "Validator Set and Finality"),
            ("3.2", "Nominator Pools"),
            ("3.3", "Slashing, Evidence, and Insurance"),
            ("3.4", "Anti-Concentration Policy"),
        ]),
        ("4", "Aetra Virtual Machine", [
            ("4.1", "Deploy and Execute Pipeline"),
            ("4.2", "External and Internal Messages"),
            ("4.3", "Gas, Storage Rent, and Exit Codes"),
            ("4.4", "Contract Standards"),
        ]),
        ("5", "Scalability and Network Operation", [
            ("5.1", "Execution Zones"),
            ("5.2", "Scheduling and Load Control"),
            ("5.3", "Experimental Sharding Roadmap"),
            ("5.4", "Validator Hardware Profile"),
        ]),
        ("6", "Native Economy", [
            ("6.1", "AET Supply"),
            ("6.2", "Fees, Burn, Treasury, and Rewards"),
            ("6.3", "Inflation and Validator Income"),
        ]),
        ("7", "Governance and Safety Gates", [
            ("7.1", "Config and Constitution"),
            ("7.2", "Public Testnet Criteria"),
            ("7.3", "Mainnet Criteria"),
        ]),
        ("Conclusion", "Conclusion", []),
        ("A", "The AET Coin", []),
    ]

    for num, title, subs in items:
        p = paragraph(doc, style="ContentsMajor")
        label = title if num == title else f"{num}   {title}"
        add_run(p, label)
        for subnum, subtitle in subs:
            p = paragraph(doc, style="ContentsMinor")
            add_run(p, f"{subnum}   {subtitle}")


def add_chapter_header(section, label):
    clear_section_links(section)
    section.different_first_page_header_footer = False
    hp = section.header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(2)
    add_run(hp, label, size=9.5, small_caps=True, spacing=16)
    bottom_rule(hp)
    fp = section.footer.paragraphs[0]
    fp.text = ""
    add_page_number(fp)


def add_chapter_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.03
    add_run(p, text, size=11)


def chapter_subheading(doc, text):
    p = paragraph(doc)
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, size=11.5, bold=True)


def chapter_one(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_page(section)
    add_chapter_header(section, "Chapter 1. Brief Description of Aetra Components")

    h = paragraph(doc)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(10)
    add_run(h, "1    Brief Description of Aetra Components", size=16, spacing=8)

    p = paragraph(doc)
    add_run(p, "Aetra", italic=True)
    add_run(
        p,
        " is intended to be a decentralized proof-of-stake execution network "
        "with a small native protocol core and a programmable contract layer. "
        "The design does not attempt to make the fastest possible blockchain. "
        "Instead, it chooses a more conservative target: practical "
        "decentralization, finality fast enough for ordinary applications, "
        "and operating requirements that do not exclude independent validators. "
        "This chapter gives a brief description of the principal components "
        "and design boundaries of the Aetra network.",
    )

    chapter_subheading(doc, "Aether Core")
    p = paragraph(doc)
    add_run(
        p,
        "Aether Core is the coordination layer of the network. It is the part "
        "of Aetra that should remain minimal, deterministic, and auditable. "
        "It maintains consensus-facing state, validator lifecycle rules, "
        "protocol configuration, fee accounting, storage accountability, "
        "routing commitments, and upgrade safety. Aether Core should not "
        "contain ordinary application business logic. Its purpose is to "
        "maintain the rules that all other execution must obey.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The core is built on a BFT proof-of-stake base. The intended network "
        "profile targets medium hardware, moderate block times, and finality "
        "that is fast enough for normal applications without requiring extreme "
        "node resources. These targets are engineering constraints rather than "
        "marketing claims. The chain should be faster than slow settlement "
        "layers while remaining easier to operate and verify than systems that "
        "assume very heavy validator hardware.",
    )

    chapter_subheading(doc, "Accounts and Address Policy")
    p = paragraph(doc)
    add_run(
        p,
        "Aetra contains an account and address layer for ordinary users, "
        "validators, contracts, and reserved system entities. This layer should "
        "make addresses readable for users while preserving strict validation "
        "rules for the protocol. It should prevent malformed addresses, "
        "zero-address misuse, and accidental overlap between user accounts and "
        "addresses reserved for system responsibilities.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The address model is also part of network clarity. A user-facing "
        "address should identify where value or messages are going, while "
        "system addresses should clearly represent built-in network duties. "
        "This separation reduces ambiguity for wallets, explorers, validators, "
        "and applications that need to distinguish user activity from protocol "
        "activity.",
    )

    chapter_subheading(doc, "Validator Economy")
    p = paragraph(doc)
    add_run(
        p,
        "Aetra's staking system is designed around validator self-stake, "
        "nominator pools, validator registry state, election logic, insurance, "
        "delegator protection, reputation, performance tracking, dynamic "
        "commission, and stake concentration controls. The purpose is not only "
        "to select validators, but also to reduce the chance that a small group "
        "of operators or economic entities can dominate the chain over time. "
        "A validator should earn enough to run reliable infrastructure, but "
        "the protocol should avoid making excessive yield the central reason "
        "to participate.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The intended validator set is deliberately not tiny. A very small set "
        "is easier to coordinate, easier to pressure, and easier to capture. "
        "A very large set at launch, on the other hand, can create consensus "
        "overhead, weak operators, and difficult synchronization. Aetra aims "
        "for a middle path where the validator set grows only when operator "
        "readiness, network stability, and finality evidence justify the "
        "increase.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "Nominator pools are included so that ordinary users do not have to "
        "manually choose a validator in order to participate in staking. A pool "
        "can aggregate stake, distribute shares, track reward indexes, process "
        "unbonding requests, and allocate stake across validators according to "
        "policy. This reduces user complexity while still allowing the protocol "
        "to apply native slashing, validator-set rules, and concentration "
        "limits at the underlying consensus layer.",
    )

    chapter_subheading(doc, "Anti-Concentration Policy")
    p = paragraph(doc)
    add_run(
        p,
        "A central design rule is that large stake should not automatically "
        "compound into larger control. If rewards are strictly proportional "
        "to already-concentrated stake, the largest validators grow faster, "
        "delegators follow them because they appear safer, and the network "
        "slowly becomes easier to cartelize. Aetra therefore includes native "
        "stake-concentration and validator-score components. These components "
        "can measure top-N concentration, apply effective voting-power caps, "
        "warn against delegation into over-concentrated validators, and reduce "
        "the reward advantage of excess stake.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The same principle applies to commission. If validator commission is "
        "unbounded, operators can manipulate delegator behavior or race to "
        "unsustainable economics. If commission is forced too low, only large "
        "operators with outside subsidies can survive. The current architecture "
        "therefore includes commission floor, ceiling, and rate-change policy, "
        "plus performance and reputation modifiers. These rules are intended "
        "to keep validator income moderate, predictable, and tied to objective "
        "operation rather than to raw size alone.",
    )

    chapter_subheading(doc, "Aetra Virtual Machine")
    p = paragraph(doc)
    add_run(
        p,
        "The Aetra Virtual Machine, or AVM, is the native smart-contract "
        "environment. It is intended to support deploy and execute flows, "
        "external and internal messages, gas accounting, host-function limits, "
        "contract storage, receipts, events, and exit codes. AVM contracts are "
        "where user-defined logic should live. A developer should be able to "
        "create new on-chain entities, define their behavior, receive messages, "
        "send messages, store state, and compose higher-level applications "
        "without changing the protocol core.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "This boundary is important. Protocol safety remains native; product "
        "and application logic moves to contracts. Staking, slashing, validator "
        "election, minting, burning, fee collection, treasury accounting, "
        "storage rent, and system configuration affect chain safety and supply "
        "correctness, so they remain built into the chain. User-defined "
        "contract logic can then evolve faster without forcing every application "
        "design into the core consensus surface.",
    )

    chapter_subheading(doc, "Storage Rent and State Accountability")
    p = paragraph(doc)
    add_run(
        p,
        "Aetra treats persistent state as a resource that must be accounted for. "
        "Contracts, account records, pool shares, reward indexes, unbonding "
        "records, and other long-lived state objects create storage pressure "
        "for every full node. Storage rent is therefore a protocol-level safety "
        "mechanism. Contracts with unpaid state costs may become frozen or "
        "limited according to explicit rules, while consensus-critical state "
        "must remain recoverable and governed by protocol-specific policies.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "This design helps avoid unbounded state growth. It also connects to "
        "the medium-hardware goal: if state can grow forever without cost, "
        "running a node becomes harder over time, and decentralization declines. "
        "Storage rent is not merely an economic feature; it is part of keeping "
        "the network practical to verify.",
    )

    chapter_subheading(doc, "System Entities")
    p = paragraph(doc)
    add_run(
        p,
        "Aetra uses native system entities for protocol responsibilities that "
        "should not be ordinary user contracts. These include configuration, "
        "constitutional bounds, validator election, validator registry, treasury, "
        "fee collection, minting, burning, emissions, reputation, performance, "
        "scheduling, storage rent, identity root, bridge coordination, and "
        "future scaling coordination. Some of these responsibilities are part "
        "of the immediate network core, while others are gated by testnet "
        "evidence and audit requirements.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "The distinction between system entities and smart contracts gives the "
        "network a clearer trust model. Built-in entities protect consensus, "
        "supply, validator power, and chain configuration. AVM contracts provide "
        "programmability for applications. This separation is intended to make "
        "Aetra flexible without making the core protocol harder to audit.",
    )

    chapter_subheading(doc, "Zones, Routing, and Future Scaling")
    p = paragraph(doc)
    add_run(
        p,
        "Aetra's architecture includes the concept of execution zones, routing, "
        "load tracking, scheduling, and future shard coordination. These ideas "
        "point toward a future where the network can increase throughput by "
        "separating execution domains and routing messages deterministically. "
        "The early network, however, should not claim production sharding until "
        "simulator coverage, adversarial tests, export/import behavior, long-run "
        "testnet evidence, and independent audit are complete.",
    )
    p = paragraph(doc)
    add_run(
        p,
        "In the intended path, Aetra starts as a simpler BFT L1 with AVM support "
        "and strong validator economics. It can then add more advanced scheduling, "
        "zone commitments, cross-zone messages, and shard coordination only when "
        "those mechanisms are deterministic and safe. This keeps the first public "
        "network understandable while preserving a path to higher throughput.",
    )


def main():
    doc = Document()
    set_page(doc.sections[0])
    configure_styles(doc)
    set_running_header(doc.sections[0], "Introduction")
    title_page(doc)
    contents_page(doc)
    chapter_one(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
