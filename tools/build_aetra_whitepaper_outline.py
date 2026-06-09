from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("docs/whitepaper/Aetra_Whitepaper_Outline_Draft.docx")


def set_page_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.38)
    section.right_margin = Inches(1.38)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_bottom_border(paragraph, color="6A6A6A", size="6"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def set_header_footer(section, label):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(label.upper())
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.font.small_caps = True
    run.font.letter_spacing = 1
    set_bottom_border(paragraph)

    footer = section.footer
    f = footer.paragraphs[0]
    add_page_number(f)
    for r in f.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = None

    h1 = styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = False
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)

    h3 = styles["Heading 3"]
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(3)

    if "AbstractBody" not in styles:
        s = styles.add_style("AbstractBody", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = "Times New Roman"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        s.font.size = Pt(10.5)
        s.paragraph_format.left_indent = Inches(0.55)
        s.paragraph_format.right_indent = Inches(0.55)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        s.paragraph_format.line_spacing = 1.05
        s.paragraph_format.space_after = Pt(4)

    if "ContentsLine" not in styles:
        s = styles.add_style("ContentsLine", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = "Times New Roman"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        s.font.size = Pt(11)
        s.paragraph_format.space_after = Pt(3)
        s.paragraph_format.tab_stops.add_tab_stop(Inches(5.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    if "ContentsChapter" not in styles:
        s = styles.add_style("ContentsChapter", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["ContentsLine"]
        s.font.bold = True
        s.font.size = Pt(12)
        s.paragraph_format.space_before = Pt(6)

    if "ContentsSub" not in styles:
        s = styles.add_style("ContentsSub", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["ContentsLine"]
        s.font.size = Pt(10.5)
        s.paragraph_format.left_indent = Inches(0.28)
        s.paragraph_format.tab_stops.add_tab_stop(Inches(5.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    if "Note" not in styles:
        s = styles.add_style("Note", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = "Times New Roman"
        s.font.size = Pt(10)
        s.paragraph_format.left_indent = Inches(0.28)
        s.paragraph_format.right_indent = Inches(0.28)
        s.paragraph_format.space_before = Pt(4)
        s.paragraph_format.space_after = Pt(8)


def p(doc, text="", style=None, align=None):
    par = doc.add_paragraph(text, style=style)
    if align is not None:
        par.alignment = align
    return par


def run(paragraph, text, *, italic=False, bold=False, size=None, small_caps=False):
    r = paragraph.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.italic = italic
    r.bold = bold
    r.font.small_caps = small_caps
    if size:
        r.font.size = Pt(size)
    return r


def add_centered_title(doc):
    for _ in range(4):
        p(doc, "")
    title = p(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(title, "Aetra", size=21)
    subtitle = p(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(subtitle, "a decentralized execution network", size=12)
    author = p(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(author, "by Daniil Shcherbakov", size=12)
    date = p(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(date, "July 8, 2026", size=12)
    for _ in range(2):
        p(doc, "")

    abstract_title = p(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(abstract_title, "Abstract", bold=True, size=10.5)
    p(
        doc,
        "This document is a preliminary whitepaper outline for Aetra, a "
        "medium-hardware, proof-of-stake L1 designed around moderate finality, "
        "anti-concentration validator economics, native system entities, and "
        "Aetra Virtual Machine (AVM) smart contracts. It is not intended to "
        "claim production readiness. Its purpose is to define the areas that "
        "the full whitepaper should cover before public testnet and mainnet.",
        style="AbstractBody",
    )
    p(
        doc,
        "The design emphasizes trust, decentralization, verifiable state "
        "correctness, storage accountability, and controlled scalability rather "
        "than pursuing maximum throughput at any cost.",
        style="AbstractBody",
    )
    for _ in range(2):
        p(doc, "")
    intro = p(doc)
    run(intro, "Introduction", size=15, bold=True)
    p(
        doc,
        "Aetra is proposed as a secure and moderately fast blockchain network "
        "with native staking infrastructure, bounded validator power, and a "
        "contract runtime built around AVM. The target is a network that can be "
        "operated by independent validators on medium hardware while preserving "
        "room for future throughput scaling through deterministic execution "
        "zones, schedulers, and experimental shard coordination.",
    )
    p(
        doc,
        "The first public versions of this paper should remain precise about "
        "implementation status. Components such as storage rent, validator "
        "election, nominator pools, AVM execution, and sharding coordination "
        "must be presented as production features only after their tests, "
        "invariants, export/import behavior, and long-running testnet evidence "
        "are complete.",
    )


contents = [
    ("1", "Brief Description of Aetra Components", "3", [
        ("1.1", "Aether Core", "3"),
        ("1.2", "Native Accounts and Address Policy", "3"),
        ("1.3", "Validator Economy", "4"),
        ("1.4", "Aetra Virtual Machine", "4"),
        ("1.5", "Execution Zones and Routing", "5"),
        ("1.6", "Storage Rent and State Accountability", "5"),
        ("1.7", "System Entities", "6"),
    ]),
    ("2", "Aetra Blockchain State", "7", [
        ("2.1", "Genesis, Export, and Import", "7"),
        ("2.2", "Messages, Transactions, and Receipts", "8"),
        ("2.3", "State Commitments and Invariants", "8"),
        ("2.4", "Upgrade and Migration Model", "9"),
    ]),
    ("3", "Consensus, Validators, and Staking", "10", [
        ("3.1", "BFT Finality and Network Profile", "10"),
        ("3.2", "Validator Registry and Election", "10"),
        ("3.3", "Nominator Pools and Liquid Staking", "11"),
        ("3.4", "Slashing, Evidence, and Insurance", "11"),
        ("3.5", "Anti-Concentration Policy", "12"),
    ]),
    ("4", "Aetra Virtual Machine", "13", [
        ("4.1", "Deploy and Execute Pipeline", "13"),
        ("4.2", "External and Internal Messages", "14"),
        ("4.3", "Gas, Host Functions, and Exit Codes", "14"),
        ("4.4", "Contract Standards", "15"),
    ]),
    ("5", "Networking, Zones, and Scalability", "16", [
        ("5.1", "Execution Zones", "16"),
        ("5.2", "Scheduler and Async Execution", "17"),
        ("5.3", "Experimental Sharding Roadmap", "17"),
        ("5.4", "Cross-Chain Registry and Bridge Hub", "18"),
    ]),
    ("6", "Native Economy", "19", [
        ("6.1", "AET and Base Denomination", "19"),
        ("6.2", "Fees, Burn, Treasury, and Rewards", "19"),
        ("6.3", "Inflation and Supply Policy", "20"),
    ]),
    ("7", "Governance and System Configuration", "21", [
        ("7.1", "Config, Constitution, and Registry", "21"),
        ("7.2", "Parameter Bounds", "22"),
        ("7.3", "Public Safety Gates", "22"),
    ]),
    ("8", "Testnet and Mainnet Readiness", "23", [
        ("8.1", "Private Validator Devnet", "23"),
        ("8.2", "Public Testnet Criteria", "24"),
        ("8.3", "Mainnet Criteria", "24"),
    ]),
    ("Conclusion", "Conclusion", "25", []),
    ("A", "The AET Coin", "26", []),
    ("B", "Implementation Checklist", "27", []),
]


def add_contents(doc):
    p(doc, "Contents", style="Heading 1")
    for num, title, page, subs in contents:
        par = p(doc, style="ContentsChapter")
        run(par, f"{num}   {title}\t{page}")
        for snum, stitle, spage in subs:
            sub = p(doc, style="ContentsSub")
            run(sub, f"{snum}   {stitle}\t{spage}")


def chapter_heading(doc, number, title):
    top = p(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(top, f"Chapter {number}. {title}", small_caps=True, size=9.5)
    set_bottom_border(top)
    par = p(doc, style="Heading 1")
    run(par, f"{number}    {title}", size=16)


def add_bullet(doc, text):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.left_indent = Inches(0.32)
    par.paragraph_format.first_line_indent = Inches(-0.18)
    par.paragraph_format.space_after = Pt(6)
    run(par, text)


def section(doc, number, title, body, bullets=None):
    par = p(doc, style="Heading 2")
    run(par, f"{number}   {title}", bold=True)
    p(doc, body)
    for item in bullets or []:
        add_bullet(doc, item)


def add_chapters(doc):
    doc.add_page_break()
    add_contents(doc)

    doc.add_page_break()
    chapter_heading(doc, "1", "Brief Description of Aetra Components")
    section(
        doc,
        "1.1",
        "Aether Core",
        "The full paper should define Aether Core as the minimal coordination "
        "layer of the network. It should own finality, validator-set changes, "
        "protocol parameters, system registries, routing commitments, and "
        "state-correctness invariants.",
        [
            "Clarify that Aether Core does not execute application business logic.",
            "List the system modules that remain native: config, constitution, scheduler, storage rent, registry, staking, evidence, treasury, mint, burn, and validator economy.",
        ],
    )
    section(
        doc,
        "1.2",
        "Native Accounts and Address Policy",
        "This section should explain user accounts, reserved system addresses, "
        "zero-address rejection, chain-id separation, and the difference between "
        "raw addresses and user-friendly encodings.",
    )
    section(
        doc,
        "1.3",
        "Validator Economy",
        "The whitepaper should present Aetra as a medium-hardware PoS network "
        "with validator self-stake, nominator pools, bounded commission, "
        "insurance, reputation, and explicit anti-concentration controls.",
    )
    section(
        doc,
        "1.4",
        "Aetra Virtual Machine",
        "AVM should be described as the native smart-contract runtime. The "
        "paper should cover deploy, execute, external messages, internal "
        "messages, gas accounting, exit codes, receipts, events, and contract "
        "standards for tokens, NFTs, domains, and exchange logic.",
    )
    section(
        doc,
        "1.5",
        "Execution Zones and Routing",
        "Execution zones should be introduced as deterministic domains for "
        "routing and isolation. In early testnet, they may be used primarily "
        "as architecture and accounting boundaries; production sharding should "
        "remain explicitly gated.",
    )
    section(
        doc,
        "1.6",
        "Storage Rent and State Accountability",
        "Aetra should treat persistent state as a scarce resource. Contracts "
        "and long-lived account records should pay storage rent or move into "
        "limited/frozen states under clearly defined rules.",
    )
    section(
        doc,
        "1.7",
        "System Entities",
        "This section should list built-in system entities and their roles: "
        "config, elector/election, treasury, fee collector, mint authority, "
        "burn authority, scheduler, storage rent, identity root, bridge hub, "
        "validator registry, reputation, performance oracle, and protection funds.",
    )

    doc.add_page_break()
    chapter_heading(doc, "2", "Aetra Blockchain State")
    section(
        doc,
        "2.1",
        "Genesis, Export, and Import",
        "The whitepaper should specify that every launch starts from a validated "
        "genesis. Export/import roundtrips, module-account balances, native "
        "supply, validator set, system addresses, and storage state must be "
        "consistent before a public network starts.",
    )
    section(
        doc,
        "2.2",
        "Messages, Transactions, and Receipts",
        "Aetra should define transaction admission, fee checks, message routing, "
        "AVM receipts, and failed-execution behavior. Receipts should expose "
        "exit code, gas used, message id, touched state roots, and events.",
    )
    section(
        doc,
        "2.3",
        "State Commitments and Invariants",
        "This section should explain app-level invariants: total supply, module "
        "accounts, staking state, validator eligibility, storage rent, system "
        "address reservations, and export/import determinism.",
    )
    section(
        doc,
        "2.4",
        "Upgrade and Migration Model",
        "Aetra should define how protocol upgrades are scheduled, how version "
        "maps are updated, how store migrations are tested, and how operators "
        "prepare for chain upgrades.",
    )

    doc.add_page_break()
    chapter_heading(doc, "3", "Consensus, Validators, and Staking")
    section(
        doc,
        "3.1",
        "BFT Finality and Network Profile",
        "The target profile is moderate speed and strong finality: blocks in "
        "the five-to-eight-second range, normal finality in the five-to-fifteen "
        "second range, and degraded finality bounded by explicit network targets.",
    )
    section(
        doc,
        "3.2",
        "Validator Registry and Election",
        "The registry should track validator identity, operator keys, consensus "
        "keys, self-stake, pool-backed stake, performance, jailing, insurance, "
        "and election eligibility.",
    )
    section(
        doc,
        "3.3",
        "Nominator Pools and Liquid Staking",
        "Ordinary users should be able to stake through official pools or "
        "staking indexes instead of manually selecting validators. The paper "
        "should explain pool shares, unbonding, reward indexes, slashing "
        "inheritance, and pool-level safety limits.",
    )
    section(
        doc,
        "3.4",
        "Slashing, Evidence, and Insurance",
        "Slashing should remain objective: double-signing, downtime, invalid "
        "state roots, invalid receipt roots, and data-unavailability claims "
        "must be backed by deterministic evidence before penalties are applied.",
    )
    section(
        doc,
        "3.5",
        "Anti-Concentration Policy",
        "The full paper should describe effective voting-power caps, reduced "
        "rewards for over-concentrated validators, entity-level monitoring, "
        "commission limits, and public concentration metrics.",
    )

    doc.add_page_break()
    chapter_heading(doc, "4", "Aetra Virtual Machine")
    section(
        doc,
        "4.1",
        "Deploy and Execute Pipeline",
        "The AVM chapter should define code storage, StateInit, contract address "
        "derivation, deployment, execution, rollback rules, and deterministic "
        "state writes.",
    )
    section(
        doc,
        "4.2",
        "External and Internal Messages",
        "Contracts should be able to receive signed external messages and "
        "contract-to-contract internal messages. The message ABI should define "
        "sender, receiver, value, fee, bounce flag, payload, nonce, and timeout.",
    )
    section(
        doc,
        "4.3",
        "Gas, Host Functions, and Exit Codes",
        "The paper should specify metered gas, host-function allowlists, memory "
        "limits, deterministic serialization, storage pricing, and built-in "
        "exit-code ranges for VM, gas, storage, message, auth, and protocol errors.",
    )
    section(
        doc,
        "4.4",
        "Contract Standards",
        "Token, NFT, domain, resolver, exchange, wallet, and registry logic "
        "should be presented as AVM standards rather than native app modules. "
        "Native modules should remain reserved for protocol-critical systems.",
    )

    doc.add_page_break()
    chapter_heading(doc, "5", "Networking, Zones, and Scalability")
    section(
        doc,
        "5.1",
        "Execution Zones",
        "Execution zones provide a way to classify state and messages by "
        "domain. The paper should separate production behavior from future "
        "research tracks so the network does not overclaim sharding readiness.",
    )
    section(
        doc,
        "5.2",
        "Scheduler and Async Execution",
        "Schedulers should coordinate timed jobs, contract queues, storage-rent "
        "collection, reward epochs, and message ordering without relying on "
        "local node timing.",
    )
    section(
        doc,
        "5.3",
        "Experimental Sharding Roadmap",
        "Sharding should be described as an R&D path until simulator tests, "
        "fuzzing, adversarial cases, validator assignment, data availability, "
        "and consensus-safety proofs are complete.",
    )
    section(
        doc,
        "5.4",
        "Cross-Chain Registry and Bridge Hub",
        "Aetra may maintain native registries for trusted chains, channels, "
        "bridge hubs, light-client references, and cross-chain routing metadata. "
        "Bridge execution should remain bounded by explicit risk controls.",
    )

    doc.add_page_break()
    chapter_heading(doc, "6", "Native Economy")
    section(
        doc,
        "6.1",
        "AET and Base Denomination",
        "The paper should define the native coin, base denomination, display "
        "denomination, precision, fee-denom policy, and the relationship between "
        "genesis balances and future supply changes.",
    )
    section(
        doc,
        "6.2",
        "Fees, Burn, Treasury, and Rewards",
        "Transaction fees should be split between burn, validator/delegator "
        "rewards, and treasury according to bounded parameters. Fee accounting "
        "must be auditable and covered by invariants.",
    )
    section(
        doc,
        "6.3",
        "Inflation and Supply Policy",
        "Inflation should be low to moderate and responsive to the bonded ratio. "
        "The paper should avoid presenting high APR as the purpose of the chain.",
    )

    doc.add_page_break()
    chapter_heading(doc, "7", "Governance and System Configuration")
    section(
        doc,
        "7.1",
        "Config, Constitution, and Registry",
        "Aetra should have native system configuration, constitutional bounds, "
        "and a registry of reserved entities. Governance may update parameters "
        "only inside explicit safety limits.",
    )
    section(
        doc,
        "7.2",
        "Parameter Bounds",
        "The whitepaper should specify bounded ranges for validator count, "
        "inflation, commission, slashing, fee split, storage rent, AVM gas, "
        "and upgrade timing.",
    )
    section(
        doc,
        "7.3",
        "Public Safety Gates",
        "Public claims should be tied to tests: genesis validation, export/import, "
        "invariants, localnet, validator operations, long-running testnet, "
        "load tests, and independent audit.",
    )

    doc.add_page_break()
    chapter_heading(doc, "8", "Testnet and Mainnet Readiness")
    section(
        doc,
        "8.1",
        "Private Validator Devnet",
        "The private devnet should validate operator setup, build reproducibility, "
        "genesis distribution, peers, restart behavior, and validator onboarding "
        "before a broad public testnet.",
    )
    section(
        doc,
        "8.2",
        "Public Testnet Criteria",
        "A public testnet should require working validator docs, stable chain-id, "
        "known slashing policy, localnet evidence, export/import tests, metrics, "
        "release binaries, and support channels.",
    )
    section(
        doc,
        "8.3",
        "Mainnet Criteria",
        "Mainnet requires a stricter bar: public testnet observation, audit, "
        "tokenomics finalization, governance safety, validator diversity, "
        "upgrade drills, and incident-response procedures.",
    )

    doc.add_page_break()
    p(doc, "Conclusion", style="Heading 1")
    p(
        doc,
        "The complete Aetra whitepaper should argue for a chain that values "
        "decentralization, verifiable correctness, and moderate operating "
        "requirements over extreme performance claims. The strongest part of "
        "the thesis is not one module, but the combination of bounded validator "
        "power, nominator pools, AVM contracts, storage accountability, and "
        "strict public-readiness gates.",
    )

    p(doc, "Appendix A    The AET Coin", style="Heading 1")
    p(
        doc,
        "This appendix should eventually define initial supply, ICO allocation "
        "policy, vesting, validator bootstrap allocation, treasury reserves, "
        "emission bounds, burn policy, and how balances are placed into mainnet "
        "genesis before launch.",
    )

    p(doc, "Appendix B    Implementation Checklist", style="Heading 1")
    add_bullet(doc, "Finalize AVM deploy/execute, gas, storage, receipts, and exit-code behavior.")
    add_bullet(doc, "Connect validator election, registry, insurance, and anti-concentration policy to actual validator-set changes.")
    add_bullet(doc, "Make nominator pools the primary user staking route and close direct user delegation if that remains the chosen policy.")
    add_bullet(doc, "Pass genesis validation, export/import roundtrip, invariants, localnet, and public testnet readiness gates.")


def main():
    doc = Document()
    section = doc.sections[0]
    set_page_geometry(section)
    section.different_first_page_header_footer = True
    configure_styles(doc)
    set_header_footer(section, "Aetra Preliminary Whitepaper")

    add_centered_title(doc)
    add_chapters(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
